from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Initialize Document
doc = Document()

# Set Default Font to Arial (Swiss Style / ATS Friendly)
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(10)

def add_section_header(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    runner = p.add_run(text.upper())
    runner.bold = True
    runner.font.name = 'Arial'
    runner.font.size = Pt(12)

# --- HEADER ---
header = doc.add_paragraph()
header.alignment = WD_ALIGN_PARAGRAPH.LEFT
name_run = header.add_run('JAMES HOOD')
name_run.bold = True
name_run.font.name = 'Arial'
name_run.font.size = Pt(24)

sub_header = doc.add_paragraph()
sub_header.paragraph_format.space_after = Pt(12)
title_run = sub_header.add_run('Senior AI Systems Architect | Verified Generalist | Full Stack Engineer')
title_run.font.name = 'Arial'
title_run.font.size = Pt(11)
title_run.bold = True

contact_info = doc.add_paragraph('St. Petersburg, FL | github.com/jameshood118 | 239-898-7833 | jameshood118@gmail.com')
contact_info.style = 'Normal'

doc.add_paragraph('_' * 90)

# --- SUMMARY ---
add_section_header('Professional Summary')
summary = doc.add_paragraph(
    '"I didn\'t pivot to AI; the industry finally built a compiler for my native language."\n\n'
    'For over 25 years, my career has been defined by a single, paradoxical methodology: I don\'t just write code... '
    'I engineer Human-to-System protocols. My expertise was forged in the trenches of crisis remediation, '
    'rapidly adopting entire tech stacks (Vue, React Native, Sitecore) over single weekends to rescue failing enterprise and military contracts.\n\n'
    'My development philosophy is grounded in the discipline of ATA Songahm Taekwondo, translating martial arts tenets into '
    'rigorous Agile protocols. Today, I formalize this "Human OS" approach by using Gemini Gems '
    'and custom AI protocols to perform Forensic System Audits, dismantle "Efficiency Traps," '
    'and build high-integrity environments where humans... and the agents that serve them... can thrive.'
)

# --- SKILLS ---
add_section_header('Technical Skills')

skills = [
    ('AI & Agentic Architecture:', 'LLM Logic Auditing, Gemini Gems (Custom Operating Protocols), Prompt Engineering for Accessibility (A11Y), Forensic Data Verification.'),
    ('Full Stack Engineering:', 'JavaScript (ES6+), TypeScript, C#, Python, React, React Native, Vue.js, Node.js, GraphQL, PostgreSQL, Supabase.'),
    ('System Integrity:', 'Web Accessibility (CPACC Standards, ARIA/TTY Remediation), Forensic Document Analysis, TDD (Jest/Cypress).'),
    ('Environment Optimization:', 'Azure GovCloud vs. Enterprise Compliance, Docker, Vercel, Git/GitLab Flow, Agile/Scrum Leadership.')
]

for title, desc in skills:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    t_run = p.add_run(title + ' ')
    t_run.bold = True
    p.add_run(desc)

# --- EXPERIENCE ---
add_section_header('Professional Experience')

# Job 1
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
run = p.add_run('WRDLNKDN | Full Stack Developer & Project Maintainer')
run.bold = True
p.add_run('\nJan 2026 - Present')

items = [
    ('Project', 'Architecting a decentralized community platform focused on dismantling "corporate theater" in hiring using React 19 and Supabase.'),
    ('Agentic Protocol', 'Implementing Safehood and Osgood-Rupert protocols to create high-integrity verification systems, ensuring candidates are evaluated on "Source Code" (actual skills) rather than "Legacy Formatting" (resumes).')
]
for label, text in items:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(f'{label}: ')
    r.bold = True
    p.add_run(text)

# Job 2
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
run = p.add_run('INDEPENDENT | Collaborative AI Developer & Forensic Analyst')
run.bold = True
p.add_run('\n2024 - Present')

items = [
    ('System Architecture', 'Developing custom Gemini AI Gems that function as localized "operating protocols," allowing for consistent, high-integrity technical strategy and content generation.'),
    ('Verification', 'Performing forensic analysis of technical documentation to identify "Hallucinations" in corporate strategy, ensuring all output meets a strict "Review/Rewind" text-first protocol.')
]
for label, text in items:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(f'{label}: ')
    r.bold = True
    p.add_run(text)

# Job 3
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
run = p.add_run('DYNAMIS, INC | Senior Software Engineer')
run.bold = True
p.add_run('\nMarch 2021 - May 2025')

items = [
    ('Rapid Deployment', 'Led the technical transition of the FEMA CBRNResponder web and mobile applications. Learned React Native over a single weekend to successfully deploy critical mobile updates and prevent contract failure.'),
    ('DoD & Enterprise Architecture', 'Engineered Full Stack React-based web solutions and DoD Operational Decision Support Software (SaaS), integrating AI algorithms into emergency management tools for global stakeholders.'),
    ('Cleared-Adjacent Operations', 'Operated in highly sensitive environments supporting NATO allies, SOCOM adjacent projects, and international health ministries without requiring traditional clearance pipelines... a testament to strict operational integrity.'),
    ('Compliance & FedRAMP', 'Navigated the complex boundaries between Azure GovCloud and Enterprise Azure, architecting scalable environments that maintained strict FedRAMP Certification and federal data sovereignty requirements.')
]
for label, text in items:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(f'{label}: ')
    r.bold = True
    p.add_run(text)

# Job 4
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
run = p.add_run('VATC, LLC | Mid-Level Web Engineer')
run.bold = True
p.add_run('\nDec 2017 - Feb 2021')

items = [
    ('Crisis Remediation', 'Architected a fully functional Vue.js and Bulma dashboard over a 48-hour period to replace a failing legacy Angular application during a critical military training exercise (Emerald Warrior).'),
    ('Synthetic Environments', 'Engineered military training exercise software by cloning major social media platforms, creating a sea of 100,000+ fake, generated posts to train soldiers in identifying bad actors.')
]
for label, text in items:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(f'{label}: ')
    r.bold = True
    p.add_run(text)

# Job 5
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
run = p.add_run('RAYMOND JAMES FINANCIAL | Web Developer (Contract)')
run.bold = True
p.add_run('\nMay 2016 - Nov 2017')

items = [
    ('Platform Adoption', 'Mastered Sitecore 7 and 8 architecture on the job to lead the migration of 3,000+ advisor websites.'),
    ('System Translation', 'Completely rewrote the technical documentation from the ground up, translating complex CMS logic so the marketing team could properly maintain accessibility and brand standards without developer intervention.')
]
for label, text in items:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(f'{label}: ')
    r.bold = True
    p.add_run(text)

# Job 6
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
run = p.add_run('CHICO\'S FAS | Front End Web Developer (Contract)')
run.bold = True
p.add_run('\nDec 2015 - April 2016')

items = [
    ('A11Y Remediation', 'Recruited specifically to halt pending ADA lawsuits across three retail brands. Remediated critical accessibility failures by implementing precise ARIA tags to resolve functional defects that caused TTY devices to silently add extra items to user carts.')
]
for label, text in items:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(f'{label}: ')
    r.bold = True
    p.add_run(text)

# --- FIELD OPS & R&D ---
add_section_header('Additional Projects & Research')
ops = [
    ('Remote Presence Protocol (2005):', 'Engineered a private, pre-YouTube residential streaming solution using analog capture cards and Windows Server. Objective: To bridge a family communication gap, allowing my wife to view real-time video of our newborn from her office in Research Park.'),
    ('Algorithmic Art & Commerce (2006-2007):', 'Pioneered procedural fractal generation using Apophysis; designs were commercially licensed by Terrapin Guitars and acquired by industry figures like Lauren Passarelli (Berklee College of Music).'),
    ('Operational Crisis Management (NovaCon 2009):', 'Voluntarily assumed a critical "Talent Liaison" role when logistical gaps emerged, managing schedules and security for Dirk Benedict, Chase Masterson, and Virginia Hey.')
]
for label, text in ops:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(f'{label} ')
    r.bold = True
    p.add_run(text)

# --- EDUCATION ---
add_section_header('Education')
edu = [
    ('Forensic Genealogical Research:', 'Ongoing audit of the Hood and McFarlin lines (1700s-Present). Demonstrates capacity for deep-dive data verification and pattern recognition.'),
    ('First-Degree Black Belt | ATA Songahm Taekwondo:', 'The source code for my Agile/Scrum discipline and focus.'),
    ('Bachelor of Science, Web Design:', 'Art Institute of Pittsburgh (2010).')
]
for label, text in edu:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(f'{label} ')
    r.bold = True
    p.add_run(text)

# Save
doc.save('James_Hood_Agentic_Resume_v6.docx')